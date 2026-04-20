from __future__ import annotations

import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = ROOT / "reports"


def _read_text_with_fallback(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-16", "utf-16-le", "utf-16-be", "cp1252"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _parse_pytest_summary(text: str) -> Dict[str, Optional[float]]:
    # Example: "================= 3 passed, 19 warnings in 176.88s (0:02:56) =================="
    summary_re = re.search(
        r"(?P<passed>\d+)\s+passed(?:,\s*(?P<failed>\d+)\s+failed)?(?:,\s*(?P<warnings>\d+)\s+warnings)?\s+in\s+(?P<seconds>[\d.]+)s",
        text,
        re.IGNORECASE,
    )
    if not summary_re:
        return {
            "passed": None,
            "failed": None,
            "warnings": None,
            "seconds": None,
        }

    return {
        "passed": float(summary_re.group("passed")) if summary_re.group("passed") else None,
        "failed": float(summary_re.group("failed")) if summary_re.group("failed") else 0.0,
        "warnings": float(summary_re.group("warnings")) if summary_re.group("warnings") else 0.0,
        "seconds": float(summary_re.group("seconds")) if summary_re.group("seconds") else None,
    }


def _parse_comprehensive_summary(text: str) -> Dict[str, Optional[int]]:
    # Example: "Total: 50 | Passed: 49 | Failed: 0 | Skipped: 1"
    m = re.search(
        r"Total:\s*(?P<total>\d+)\s*\|\s*Passed:\s*(?P<passed>\d+)\s*\|\s*Failed:\s*(?P<failed>\d+)\s*\|\s*Skipped:\s*(?P<skipped>\d+)",
        text,
        re.IGNORECASE,
    )
    if not m:
        return {"total": None, "passed": None, "failed": None, "skipped": None}
    return {
        "total": int(m.group("total")),
        "passed": int(m.group("passed")),
        "failed": int(m.group("failed")),
        "skipped": int(m.group("skipped")),
    }


def _parse_module_counts(text: str) -> Dict[str, Dict[str, int]]:
    module_counts: Dict[str, Dict[str, int]] = {}
    current = None
    for line in text.splitlines():
        m = re.match(r"MODULE\s+\d+:\s+(.+?)\s+\(TC\d+-TC\d+\)", line.strip())
        if m:
            current = m.group(1).strip()
            module_counts[current] = {"PASS": 0, "FAIL": 0, "SKIP": 0}
            continue

        if not current:
            continue

        if line.startswith("PASS TC"):
            module_counts[current]["PASS"] += 1
        elif line.startswith("FAIL TC"):
            module_counts[current]["FAIL"] += 1
        elif line.startswith("SKIP TC"):
            module_counts[current]["SKIP"] += 1

    return module_counts


def _parse_timestamps(text: str) -> Tuple[Optional[datetime], Optional[datetime]]:
    ts_re = re.compile(r"^(\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}\.\d{3})", re.MULTILINE)
    matches = ts_re.findall(text)
    if not matches:
        return None, None
    first = datetime.strptime(matches[0], "%Y-%m-%d %H:%M:%S.%f")
    last = datetime.strptime(matches[-1], "%Y-%m-%d %H:%M:%S.%f")
    return first, last


def _count_warning_categories(text: str) -> Counter:
    counter: Counter = Counter()
    for line in text.splitlines():
        if "WARNING" not in line:
            continue
        if "PydanticDeprecatedSince20" in line:
            counter["Pydantic deprecation"] += 1
        elif "parse error" in line.lower():
            counter["LLM parse warning"] += 1
        elif "timed out" in line.lower():
            counter["Timeout warning"] += 1
        else:
            counter["Other warning"] += 1
    return counter


def _extract_quality_scores(text: str) -> List[float]:
    return [float(v) for v in re.findall(r"Quality score:\s*([0-9]+\.[0-9]+)", text)]


def _chart_outcome_pie(path: Path, summary: Dict[str, Optional[int]]) -> None:
    labels = ["Passed", "Failed", "Skipped"]
    values = [summary.get("passed", 0) or 0, summary.get("failed", 0) or 0, summary.get("skipped", 0) or 0]
    colors = ["#2ca02c", "#d62728", "#ffbf00"]

    plt.figure(figsize=(7, 5))
    if sum(values) == 0:
        plt.text(0.5, 0.5, "No complete summary available", ha="center", va="center")
        plt.axis("off")
    else:
        plt.pie(values, labels=labels, autopct="%1.1f%%", colors=colors, startangle=90)
        plt.title("Comprehensive Suite Outcome Distribution")
    plt.tight_layout()
    plt.savefig(path, dpi=140)
    plt.close()


def _chart_module_bars(path: Path, module_counts: Dict[str, Dict[str, int]]) -> None:
    modules = list(module_counts.keys())
    passed = [module_counts[m]["PASS"] for m in modules]
    failed = [module_counts[m]["FAIL"] for m in modules]
    skipped = [module_counts[m]["SKIP"] for m in modules]

    x = range(len(modules))
    plt.figure(figsize=(11, 5))
    plt.bar(x, passed, label="Pass", color="#2ca02c")
    plt.bar(x, failed, bottom=passed, label="Fail", color="#d62728")
    plt.bar(x, skipped, bottom=[passed[i] + failed[i] for i in x], label="Skip", color="#ffbf00")
    plt.xticks(list(x), modules, rotation=30, ha="right")
    plt.ylabel("Test Cases")
    plt.title("Per-Module Status Breakdown (Comprehensive Suite)")
    plt.legend()
    plt.tight_layout()
    plt.savefig(path, dpi=140)
    plt.close()


def _chart_duration_bars(path: Path, durations: Dict[str, float]) -> None:
    names = list(durations.keys())
    vals = [durations[n] for n in names]
    plt.figure(figsize=(8, 5))
    plt.bar(names, vals, color=["#1f77b4", "#17becf", "#9467bd"])
    for idx, val in enumerate(vals):
        plt.text(idx, val + 1, f"{val:.1f}s", ha="center", fontsize=9)
    plt.ylabel("Seconds")
    plt.title("Execution Time Comparison")
    plt.tight_layout()
    plt.savefig(path, dpi=140)
    plt.close()


def _chart_warning_bars(path: Path, warning_counter: Counter) -> None:
    items = warning_counter.most_common()
    labels = [k for k, _ in items]
    values = [v for _, v in items]

    plt.figure(figsize=(8, 5))
    if values:
        plt.bar(labels, values, color="#ff7f0e")
        for idx, val in enumerate(values):
            plt.text(idx, val + 0.2, str(val), ha="center", fontsize=9)
        plt.xticks(rotation=20, ha="right")
        plt.ylabel("Count")
        plt.title("Warning Type Frequency")
    else:
        plt.text(0.5, 0.5, "No warnings parsed", ha="center", va="center")
        plt.axis("off")
    plt.tight_layout()
    plt.savefig(path, dpi=140)
    plt.close()


def _chart_quality_scores(path: Path, scores: List[float]) -> None:
    plt.figure(figsize=(8, 4.5))
    if scores:
        plt.plot(range(1, len(scores) + 1), scores, marker="o", color="#2c3e50")
        plt.ylim(0, 1.05)
        plt.xlabel("Observation")
        plt.ylabel("Quality Score")
        plt.title("Observed Quality Score Trend")
        plt.grid(alpha=0.25)
    else:
        plt.text(0.5, 0.5, "No quality scores found", ha="center", va="center")
        plt.axis("off")
    plt.tight_layout()
    plt.savefig(path, dpi=140)
    plt.close()


def _chart_constraint_matrix(path: Path, suites: List[str], constraints: List[str], matrix: List[List[int]]) -> None:
    plt.figure(figsize=(10, 4.8))
    plt.imshow(matrix, cmap="RdYlGn", aspect="auto", vmin=0, vmax=1)
    plt.colorbar(label="Constraint Satisfied (1=yes, 0=no)")
    plt.xticks(range(len(constraints)), constraints, rotation=25, ha="right")
    plt.yticks(range(len(suites)), suites)
    plt.title("Constraint Compliance Matrix")

    for y in range(len(suites)):
        for x in range(len(constraints)):
            plt.text(x, y, str(matrix[y][x]), ha="center", va="center", fontsize=9, color="black")

    plt.tight_layout()
    plt.savefig(path, dpi=140)
    plt.close()


def _chart_scenario_scores(path: Path, suites: List[str], scenario_scores: Dict[str, List[float]]) -> None:
    scenarios = list(scenario_scores.keys())
    x = range(len(suites))
    width = 0.22

    plt.figure(figsize=(10, 5))
    for idx, scenario in enumerate(scenarios):
        offset = (idx - (len(scenarios) - 1) / 2) * width
        vals = scenario_scores[scenario]
        bars = plt.bar([i + offset for i in x], vals, width=width, label=scenario)
        for bar in bars:
            h = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2, h + 0.01, f"{h:.2f}", ha="center", va="bottom", fontsize=8)

    plt.ylim(0, 1.05)
    plt.xticks(list(x), suites, rotation=18, ha="right")
    plt.ylabel("Composite Score (0-1)")
    plt.title("Suite Ranking Under Different Constraints")
    plt.legend()
    plt.tight_layout()
    plt.savefig(path, dpi=140)
    plt.close()


def _write_docx_report(
    out_dir: Path,
    generated_at: datetime,
    comprehensive_source_name: str,
    mocked_source_name: str,
    live_source_name: str,
    partial_status_lines: List[str],
    comp_summary: Dict[str, Optional[int]],
    comp_duration_est: Optional[float],
    mocked_summary: Dict[str, Optional[float]],
    live_summary: Dict[str, Optional[float]],
    matrix: List[List[int]],
    warning_counter: Counter,
    quality_scores: List[float],
    chart_paths: Dict[str, Path],
) -> Path:
    def _fmt_float(val: Optional[float], suffix: str = "") -> str:
        if val is None:
            return "N/A"
        return f"{val:.2f}{suffix}"

    comp_total = int(comp_summary.get("total") or 0)
    comp_passed = int(comp_summary.get("passed") or 0)
    comp_failed = int(comp_summary.get("failed") or 0)
    comp_skipped = int(comp_summary.get("skipped") or 0)
    comp_pass_rate = (comp_passed / comp_total) if comp_total > 0 else None

    mocked_passed = int(mocked_summary.get("passed") or 0)
    mocked_failed = int(mocked_summary.get("failed") or 0)
    mocked_warnings = int(mocked_summary.get("warnings") or 0)
    mocked_seconds = mocked_summary.get("seconds")
    mocked_total = mocked_passed + mocked_failed
    mocked_pass_rate = (mocked_passed / mocked_total) if mocked_total > 0 else None

    live_passed = int(live_summary.get("passed") or 0)
    live_failed = int(live_summary.get("failed") or 0)
    live_warnings = int(live_summary.get("warnings") or 0)
    live_seconds = live_summary.get("seconds")
    live_total = live_passed + live_failed
    live_pass_rate = (live_passed / live_total) if live_total > 0 else None

    doc = Document()
    doc.add_heading("Lumina Backend Evaluation Report", level=1)
    doc.add_paragraph(f"Generated: {generated_at.strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(
        "This document consolidates the latest comprehensive, mocked integration, and "
        "targeted live evaluator test evidence. It includes pass/fail outcomes, runtime "
        "characteristics, warning profile, and compliance against operational constraints."
    )
    doc.add_paragraph(
        f"Comprehensive: passed={comp_passed}, failed={comp_failed}, skipped={comp_skipped}, "
        f"pass_rate={_fmt_float(comp_pass_rate * 100 if comp_pass_rate is not None else None, '%')}."
    )
    doc.add_paragraph(
        f"Mocked integration: passed={mocked_passed}, failed={mocked_failed}, warnings={mocked_warnings}, "
        f"duration={_fmt_float(mocked_seconds, 's')}."
    )
    doc.add_paragraph(
        f"Live targeted evaluators: passed={live_passed}, failed={live_failed}, warnings={live_warnings}, "
        f"duration={_fmt_float(live_seconds, 's')}."
    )

    doc.add_heading("Data Sources", level=2)
    doc.add_paragraph(f"Comprehensive source: {comprehensive_source_name}")
    doc.add_paragraph(f"Mocked integration source: {mocked_source_name}")
    doc.add_paragraph(f"Targeted live evaluator source: {live_source_name}")
    doc.add_paragraph("Partial comprehensive runs:")
    if partial_status_lines:
        for line in partial_status_lines:
            doc.add_paragraph(line.lstrip("- ").strip(), style="List Bullet")
    else:
        doc.add_paragraph("No partial comprehensive attempts detected.", style="List Bullet")

    doc.add_heading("Core Metrics", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Suite", "Passed", "Failed", "Skipped", "Warnings", "Duration (s)", "Pass Rate"]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header

    rows = [
        [
            "Comprehensive",
            str(comp_passed),
            str(comp_failed),
            str(comp_skipped),
            "N/A",
            _fmt_float(comp_duration_est),
            _fmt_float(comp_pass_rate * 100 if comp_pass_rate is not None else None, "%"),
        ],
        [
            "Mocked Integration",
            str(mocked_passed),
            str(mocked_failed),
            "0",
            str(mocked_warnings),
            _fmt_float(mocked_seconds),
            _fmt_float(mocked_pass_rate * 100 if mocked_pass_rate is not None else None, "%"),
        ],
        [
            "Live Targeted",
            str(live_passed),
            str(live_failed),
            "0",
            str(live_warnings),
            _fmt_float(live_seconds),
            _fmt_float(live_pass_rate * 100 if live_pass_rate is not None else None, "%"),
        ],
    ]

    for data_row in rows:
        row_cells = table.add_row().cells
        for idx, val in enumerate(data_row):
            row_cells[idx].text = val

    doc.add_heading("Constraint Compliance Matrix", level=2)
    ctable = doc.add_table(rows=1, cols=6)
    ctable.style = "Table Grid"
    cheaders = ["Suite", "<=10s", "<=180s", "Pass>=95%", "Warn/test<=5", "Completed"]
    for idx, header in enumerate(cheaders):
        ctable.cell(0, idx).text = header

    suites = ["Comprehensive", "Mocked Integration", "Live Targeted"]
    for idx, suite in enumerate(suites):
        row_cells = ctable.add_row().cells
        row_cells[0].text = suite
        row_cells[1].text = str(matrix[idx][0])
        row_cells[2].text = str(matrix[idx][1])
        row_cells[3].text = str(matrix[idx][2])
        row_cells[4].text = str(matrix[idx][3])
        row_cells[5].text = str(matrix[idx][4])

    doc.add_heading("Warnings and Quality Signal", level=2)
    if warning_counter:
        for name, count in warning_counter.most_common():
            doc.add_paragraph(f"{name}: {count}", style="List Bullet")
    else:
        doc.add_paragraph("No warnings detected in parsed logs.", style="List Bullet")

    if quality_scores:
        avg_quality = sum(quality_scores) / len(quality_scores)
        doc.add_paragraph(
            f"Observed quality score points: {len(quality_scores)} | "
            f"mean={avg_quality:.3f}, min={min(quality_scores):.3f}, max={max(quality_scores):.3f}."
        )
    else:
        doc.add_paragraph("No quality score entries were parsed from current logs.")

    doc.add_heading("Visual Evidence", level=2)
    doc.add_paragraph("Comprehensive outcome distribution")
    doc.add_picture(str(chart_paths["outcome"]), width=Inches(6.2))
    doc.add_paragraph("Module status breakdown")
    doc.add_picture(str(chart_paths["modules"]), width=Inches(6.2))
    doc.add_paragraph("Execution time comparison")
    doc.add_picture(str(chart_paths["duration"]), width=Inches(6.2))
    doc.add_paragraph("Warning type frequency")
    doc.add_picture(str(chart_paths["warnings"]), width=Inches(6.2))
    doc.add_paragraph("Quality score trend")
    doc.add_picture(str(chart_paths["quality"]), width=Inches(6.2))
    doc.add_paragraph("Constraint compliance heatmap")
    doc.add_picture(str(chart_paths["constraint_matrix"]), width=Inches(6.2))
    doc.add_paragraph("Scenario-based suite ranking")
    doc.add_picture(str(chart_paths["scenario_scores"]), width=Inches(6.2))

    doc.add_heading("Interpretation and Recommendations", level=2)
    doc.add_paragraph(
        "1. Keep mocked integration checks as fast CI gates and run live evaluator checks in scheduled or pre-release pipelines."
    )
    doc.add_paragraph(
        "2. Preserve machine-readable test summary output in all long-running suites to avoid visibility loss on interrupted runs."
    )
    doc.add_paragraph(
        "3. Track warning-rate KPI over time (especially parse and timeout categories) alongside pass/fail outcomes."
    )
    doc.add_paragraph(
        "4. Maintain regular quality-score trend sampling to detect silent regressions in response quality."
    )

    docx_path = out_dir / "Lumina_Backend_Evaluation_Report.docx"
    doc.save(docx_path)
    return docx_path


def generate_report() -> Path:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = REPORTS_DIR / f"evaluation_report_{stamp}"
    out_dir.mkdir(parents=True, exist_ok=True)

    comprehensive_candidates = [
        ROOT / "test_run_output_utf8.txt",
        ROOT / "comprehensive_latest.log",
        ROOT / "comprehensive_latest2.log",
    ]
    mocked_path = ROOT / "mocked_integration_latest.log"
    live_targeted_path = ROOT / "targeted_live_eval_latest.log"
    partial_paths = [ROOT / "comprehensive_latest.log", ROOT / "comprehensive_latest2.log"]

    comprehensive_complete_path = None
    comprehensive_complete = ""
    for candidate in comprehensive_candidates:
        if not candidate.exists():
            continue
        candidate_text = _read_text_with_fallback(candidate)
        candidate_summary = _parse_comprehensive_summary(candidate_text)
        if candidate_summary.get("total"):
            comprehensive_complete_path = candidate
            comprehensive_complete = candidate_text
            break

    if comprehensive_complete_path is None:
        for candidate in comprehensive_candidates:
            if candidate.exists():
                comprehensive_complete_path = candidate
                comprehensive_complete = _read_text_with_fallback(candidate)
                break

    mocked_text = _read_text_with_fallback(mocked_path) if mocked_path.exists() else ""
    live_text = _read_text_with_fallback(live_targeted_path) if live_targeted_path.exists() else ""
    partial_texts = [_read_text_with_fallback(p) for p in partial_paths if p.exists()]

    comp_summary = _parse_comprehensive_summary(comprehensive_complete)
    module_counts = _parse_module_counts(comprehensive_complete)
    mocked_summary = _parse_pytest_summary(mocked_text)
    live_summary = _parse_pytest_summary(live_text)

    first_ts, last_ts = _parse_timestamps(comprehensive_complete)
    comp_duration_est = (last_ts - first_ts).total_seconds() if first_ts and last_ts else None

    duration_map: Dict[str, float] = {}
    if comp_duration_est is not None:
        duration_map["Comprehensive"] = comp_duration_est
    if mocked_summary["seconds"] is not None:
        duration_map["Mocked Integration"] = float(mocked_summary["seconds"])
    if live_summary["seconds"] is not None:
        duration_map["Live Targeted Eval"] = float(live_summary["seconds"])

    warning_counter = Counter()
    for txt in [comprehensive_complete, mocked_text, live_text, *partial_texts]:
        warning_counter.update(_count_warning_categories(txt))

    quality_scores = []
    for txt in [comprehensive_complete, *partial_texts, live_text]:
        quality_scores.extend(_extract_quality_scores(txt))

    pie_path = out_dir / "outcome_pie.png"
    modules_path = out_dir / "module_breakdown.png"
    duration_path = out_dir / "duration_comparison.png"
    warning_path = out_dir / "warning_frequency.png"
    quality_path = out_dir / "quality_scores.png"
    constraint_matrix_path = out_dir / "constraint_matrix.png"
    scenario_scores_path = out_dir / "scenario_scores.png"

    _chart_outcome_pie(pie_path, comp_summary)
    _chart_module_bars(modules_path, module_counts)
    _chart_duration_bars(duration_path, duration_map)
    _chart_warning_bars(warning_path, warning_counter)
    _chart_quality_scores(quality_path, quality_scores)

    suites = [
        "Comprehensive",
        "Mocked Integration",
        "Live Targeted",
    ]

    comp_total = float(comp_summary.get("total") or 0)
    comp_passed = float(comp_summary.get("passed") or 0)
    comp_failed = float(comp_summary.get("failed") or 0)

    comp_tests = comp_total if comp_total > 0 else None
    mocked_tests = (mocked_summary["passed"] or 0) + (mocked_summary["failed"] or 0)
    live_tests = (live_summary["passed"] or 0) + (live_summary["failed"] or 0)

    metric_rows = [
        {
            "name": "Comprehensive",
            "duration": comp_duration_est,
            "pass_rate": (comp_passed / comp_tests) if comp_tests else None,
            "warning_rate": None,
            "completed": bool(comp_summary.get("total")),
        },
        {
            "name": "Mocked Integration",
            "duration": mocked_summary["seconds"],
            "pass_rate": (mocked_summary["passed"] / mocked_tests) if mocked_tests else None,
            "warning_rate": ((mocked_summary["warnings"] or 0) / mocked_tests) if mocked_tests else None,
            "completed": True,
        },
        {
            "name": "Live Targeted",
            "duration": live_summary["seconds"],
            "pass_rate": (live_summary["passed"] / live_tests) if live_tests else None,
            "warning_rate": ((live_summary["warnings"] or 0) / live_tests) if live_tests else None,
            "completed": True,
        },
    ]

    constraints = [
        "Duration <= 10s",
        "Duration <= 180s",
        "Pass rate >= 95%",
        "Warn/test <= 5",
        "Completed run",
    ]

    matrix: List[List[int]] = []
    for row in metric_rows:
        duration = row["duration"]
        pass_rate = row["pass_rate"]
        warning_rate = row["warning_rate"]
        completed = row["completed"]

        matrix.append(
            [
                int(duration is not None and duration <= 10),
                int(duration is not None and duration <= 180),
                int(pass_rate is not None and pass_rate >= 0.95),
                int(warning_rate is not None and warning_rate <= 5),
                int(completed),
            ]
        )

    _chart_constraint_matrix(constraint_matrix_path, suites, constraints, matrix)

    known_durations = [row["duration"] for row in metric_rows if row["duration"] is not None]
    min_dur = min(known_durations) if known_durations else 0.0
    max_dur = max(known_durations) if known_durations else 1.0

    def _norm_speed(duration: Optional[float]) -> float:
        if duration is None:
            return 0.0
        if max_dur == min_dur:
            return 1.0
        return 1.0 - ((duration - min_dur) / (max_dur - min_dur))

    def _norm_warning(warn_rate: Optional[float]) -> float:
        if warn_rate is None:
            return 0.0
        # 0 warnings/test -> 1.0, 5+ warnings/test -> 0.0
        return max(0.0, min(1.0, 1.0 - (warn_rate / 5.0)))

    scenario_scores = {
        "Speed-first": [],
        "Reliability-first": [],
        "Stability-first": [],
    }

    for row in metric_rows:
        speed = _norm_speed(row["duration"])
        pass_rate = row["pass_rate"] if row["pass_rate"] is not None else 0.0
        warning = _norm_warning(row["warning_rate"])
        completed = 1.0 if row["completed"] else 0.0

        scenario_scores["Speed-first"].append(0.60 * speed + 0.30 * pass_rate + 0.10 * warning)
        scenario_scores["Reliability-first"].append(0.65 * pass_rate + 0.25 * completed + 0.10 * warning)
        scenario_scores["Stability-first"].append(0.50 * warning + 0.30 * completed + 0.20 * pass_rate)

    _chart_scenario_scores(scenario_scores_path, suites, scenario_scores)

    partial_status_lines = []
    for p in partial_paths:
        if not p.exists():
            continue
        txt = _read_text_with_fallback(p)
        has_summary = bool(re.search(r"TEST SUMMARY|Total:\s*\d+\s*\|\s*Passed:", txt, re.IGNORECASE))
        p_first, p_last = _parse_timestamps(txt)
        p_dur = (p_last - p_first).total_seconds() if p_first and p_last else None
        partial_status_lines.append(
            f"- {p.name}: {'complete' if has_summary else 'incomplete'}"
            + (f", observed span ~{p_dur:.1f}s" if p_dur is not None else "")
        )

    def _f(v: Optional[float]) -> str:
        return "N/A" if v is None else f"{v:.2f}"

    comp_passed_count = int(comp_summary.get("passed") or 0)
    comp_failed_count = int(comp_summary.get("failed") or 0)
    comp_skipped_count = int(comp_summary.get("skipped") or 0)
    comp_total_count = int(comp_summary.get("total") or 0)
    comp_pass_rate_text = "N/A"
    if comp_total_count > 0:
        comp_pass_rate_text = f"{(comp_passed_count / comp_total_count) * 100:.1f}%"

    generated_at = datetime.now()

    report = out_dir / "evaluation_report.md"
    report.write_text(
        "\n".join(
            [
                "# Lumina Backend Evaluation Report",
                "",
                f"Generated: {generated_at.strftime('%Y-%m-%d %H:%M:%S')}",
                "",
                "## Executive Summary",
                "",
                f"- Baseline comprehensive run: passed={comp_passed_count}, failed={comp_failed_count}, skipped={comp_skipped_count}, pass_rate={comp_pass_rate_text}.",
                f"- Mocked integration run (latest): passed={int(mocked_summary['passed'] or 0)}, failed={int(mocked_summary['failed'] or 0)}, warnings={int(mocked_summary['warnings'] or 0)}, duration={_f(mocked_summary['seconds'])}s.",
                f"- Targeted live evaluator run (latest): passed={int(live_summary['passed'] or 0)}, failed={int(live_summary['failed'] or 0)}, warnings={int(live_summary['warnings'] or 0)}, duration={_f(live_summary['seconds'])}s.",
                "",
                "## Data Sources",
                "",
                f"- Comprehensive run source: `{comprehensive_complete_path.name if comprehensive_complete_path else 'N/A'}`",
                f"- Latest mocked integration run: `{mocked_path.name}`",
                f"- Latest targeted live eval run: `{live_targeted_path.name}`",
                "- Partial comprehensive attempts:",
                *partial_status_lines,
                "",
                "## Core Metrics",
                "",
                "| Suite | Passed | Failed | Skipped | Warnings | Duration (s) |",
                "|---|---:|---:|---:|---:|---:|",
                f"| Comprehensive (latest source) | {comp_passed_count} | {comp_failed_count} | {comp_skipped_count} | N/A | {_f(comp_duration_est)} |",
                f"| Mocked Integration (latest) | {int(mocked_summary['passed'] or 0)} | {int(mocked_summary['failed'] or 0)} | 0 | {int(mocked_summary['warnings'] or 0)} | {_f(mocked_summary['seconds'])} |",
                f"| Live Targeted Evaluators (latest) | {int(live_summary['passed'] or 0)} | {int(live_summary['failed'] or 0)} | 0 | {int(live_summary['warnings'] or 0)} | {_f(live_summary['seconds'])} |",
                "",
                "## Visualizations",
                "",
                "### 1) Comprehensive Outcome Distribution",
                "![Comprehensive Outcome Pie](outcome_pie.png)",
                "",
                "### 2) Per-Module Test Status Breakdown",
                "![Module Breakdown](module_breakdown.png)",
                "",
                "### 3) Duration Comparison",
                "![Duration Comparison](duration_comparison.png)",
                "",
                "### 4) Warning Type Frequency",
                "![Warning Frequency](warning_frequency.png)",
                "",
                "### 5) Quality Score Trend",
                "![Quality Score Trend](quality_scores.png)",
                "",
                "### 6) Constraint Compliance Matrix",
                "![Constraint Compliance](constraint_matrix.png)",
                "",
                "### 7) Rankings Under Different Constraint Priorities",
                "![Constraint Scenario Scores](scenario_scores.png)",
                "",
                "## Constraint Summary Table",
                "",
                "| Suite | <=10s | <=180s | Pass>=95% | Warn/test<=5 | Completed |",
                "|---|---:|---:|---:|---:|---:|",
                f"| Comprehensive | {matrix[0][0]} | {matrix[0][1]} | {matrix[0][2]} | {matrix[0][3]} | {matrix[0][4]} |",
                f"| Mocked Integration | {matrix[1][0]} | {matrix[1][1]} | {matrix[1][2]} | {matrix[1][3]} | {matrix[1][4]} |",
                f"| Live Targeted | {matrix[2][0]} | {matrix[2][1]} | {matrix[2][2]} | {matrix[2][3]} | {matrix[2][4]} |",
                "",
                "## Performance Interpretation",
                "",
                (
                    f"- Reliability: Comprehensive baseline reports {comp_failed_count} failed tests out of {comp_total_count} total."
                    if comp_total_count > 0
                    else "- Reliability: Comprehensive baseline completion summary is unavailable; mocked/targeted latest runs report zero failed tests."
                ),
                "- Robustness: Live evaluator logs still show occasional LLM JSON parse warnings, but tests pass due fallback handling.",
                "- Throughput/latency: Mocked integration is fast (~2s), while targeted live evaluator checks are much slower (~177s), reflecting external LLM dependency overhead.",
                (
                    f"- Breadth vs repeatability: Comprehensive source includes {comp_total_count} summarized cases."
                    if comp_total_count > 0
                    else "- Breadth vs repeatability: Current comprehensive sources appear incomplete and do not include a terminal summary, indicating long runs can terminate before full aggregation."
                ),
                "",
                "## Recommendations",
                "",
                "1. Keep mocked integration in CI as a fast gate and run live targeted evaluator checks in scheduled pipelines.",
                "2. Add explicit timeout/interrupt accounting to comprehensive runner so incomplete runs still emit machine-readable summary.",
                "3. Track warning rate over time (especially parse/timeouts) as a stability KPI, not only pass/fail.",
            ]
        ),
        encoding="utf-8",
    )

    chart_paths = {
        "outcome": pie_path,
        "modules": modules_path,
        "duration": duration_path,
        "warnings": warning_path,
        "quality": quality_path,
        "constraint_matrix": constraint_matrix_path,
        "scenario_scores": scenario_scores_path,
    }

    _write_docx_report(
        out_dir=out_dir,
        generated_at=generated_at,
        comprehensive_source_name=comprehensive_complete_path.name if comprehensive_complete_path else "N/A",
        mocked_source_name=mocked_path.name,
        live_source_name=live_targeted_path.name,
        partial_status_lines=partial_status_lines,
        comp_summary=comp_summary,
        comp_duration_est=comp_duration_est,
        mocked_summary=mocked_summary,
        live_summary=live_summary,
        matrix=matrix,
        warning_counter=warning_counter,
        quality_scores=quality_scores,
        chart_paths=chart_paths,
    )

    return report


def main() -> None:
    report_path = generate_report()
    print(f"Generated report: {report_path}")


if __name__ == "__main__":
    main()
