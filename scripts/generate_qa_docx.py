from pathlib import Path

from docx import Document
from docx.shared import Pt


def split_table_row(line: str) -> list[str]:
    text = line.strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|"):
        text = text[:-1]
    return [cell.strip() for cell in text.split("|")]


def is_table_separator(line: str) -> bool:
    stripped = line.replace("|", "").replace("-", "").replace(":", "").strip()
    return stripped == ""


def build_docx_from_markdown(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()

    in_code_block = False
    code_buffer: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_buffer = []
            else:
                para = doc.add_paragraph("\n".join(code_buffer))
                para.style = doc.styles["No Spacing"]
                for run in para.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                in_code_block = False
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|") and lines[i].endswith("|"):
                table_lines.append(lines[i])
                i += 1

            if len(table_lines) >= 2 and is_table_separator(table_lines[1]):
                headers = split_table_row(table_lines[0])
                data_rows = [split_table_row(r) for r in table_lines[2:]]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for col, text in enumerate(headers):
                    table.cell(0, col).text = text
                for row_data in data_rows:
                    row = table.add_row().cells
                    for col in range(len(headers)):
                        row[col].text = row_data[col] if col < len(row_data) else ""
            else:
                for table_line in table_lines:
                    doc.add_paragraph(table_line)
            continue

        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            i += 1
            continue

        stripped = line.strip()
        if stripped and stripped[0].isdigit() and ". " in stripped:
            marker, remainder = stripped.split(". ", 1)
            if marker.isdigit():
                doc.add_paragraph(remainder.strip(), style="List Number")
                i += 1
                continue

        if line.strip() == "---":
            doc.add_paragraph("")
            i += 1
            continue

        doc.add_paragraph(line)
        i += 1

    doc.save(docx_path)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    md_path = root / "LUMINA_QA_Testing_Document.md"
    docx_path = root / "LUMINA_QA_Testing_Document.docx"
    build_docx_from_markdown(md_path, docx_path)
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    main()
